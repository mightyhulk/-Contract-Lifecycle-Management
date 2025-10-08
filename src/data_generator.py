#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Synthetic Contract Data Generator

This script generates a synthetic dataset of contracts for testing the CLM system.
"""

import os
import random
from datetime import datetime, timedelta
from docx import Document
import io
from PIL import Image, ImageDraw, ImageFont
import pytesseract
from PyPDF2 import PdfWriter, PdfReader
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("logs/data_generator.log"),
        logging.StreamHandler()
    ]
)

logger = logging.getLogger(__name__)

# Constants
OUTPUT_DIR = "data/contracts"
CURRENT_DATE = datetime.now()
COMPANY_DATA = {
    "ABC Corp": {
        "addresses": [
            "123 Main St, New York, NY 10001", 
            "456 Broadway Ave, New York, NY 10002"  # Conflicting address
        ],
        "contact": ["john.doe@abccorp.com", "212-555-1234"],
    },
    "XYZ Industries": {
        "addresses": [
            "789 Park Ave, Chicago, IL 60601",
            "987 Lake St, Chicago, IL 60602"  # Conflicting address
        ],
        "contact": ["jane.smith@xyzind.com", "312-555-6789"],
    },
    "Tech Solutions LLC": {
        "addresses": [
            "101 Tech Blvd, San Francisco, CA 94107",
            "202 Innovation Way, San Francisco, CA 94108"  # Conflicting address
        ],
        "contact": ["info@techsolutions.com", "415-555-4321"],
    },
    "Global Services Inc": {
        "addresses": [
            "300 Corporate Dr, Boston, MA 02110",
            "400 Business Pkwy, Boston, MA 02111"  # Conflicting address
        ],
        "contact": ["contact@globalservices.com", "617-555-8765"],
    }
}

CONTRACT_TYPES = [
    "Service Agreement", 
    "Sales Contract", 
    "Partnership Agreement",
    "Maintenance Contract", 
    "Software License"
]

DEPARTMENTS = [
    "Legal", 
    "Procurement", 
    "Sales", 
    "IT", 
    "Operations"
]

def random_date(start, end):
    """Generate a random date between start and end dates."""
    delta = end - start
    int_delta = (delta.days * 24 * 60 * 60) + delta.seconds
    random_second = random.randrange(int_delta)
    return start + timedelta(seconds=random_second)

def generate_contract_dates(with_conflicts=False):
    """Generate contract dates, optionally with conflicts."""
    creation_date = random_date(
        CURRENT_DATE - timedelta(days=365*2), 
        CURRENT_DATE - timedelta(days=30)
    )
    
    # Contract duration between 1 and 3 years
    duration = random.randint(1, 3)
    
    if with_conflicts and random.random() < 0.3:
        # Create a conflict by having two different expiration dates
        expiration_date_1 = creation_date + timedelta(days=365*duration)
        expiration_date_2 = creation_date + timedelta(days=365*duration + random.randint(30, 90))
        expiration_date = random.choice([expiration_date_1, expiration_date_2])
    else:
        expiration_date = creation_date + timedelta(days=365*duration)
    
    # Set renewal date 30 days before expiration
    renewal_date = expiration_date - timedelta(days=30)
    
    return {
        "creation": creation_date,
        "expiration": expiration_date,
        "renewal": renewal_date
    }

def select_parties(conflict=False):
    """Select contracting parties."""
    parties = random.sample(list(COMPANY_DATA.keys()), 2)
    
    party_details = {}
    for party in parties:
        if conflict and random.random() < 0.3:
            # Introduce conflict in address
            address = random.choice(COMPANY_DATA[party]["addresses"])
        else:
            address = COMPANY_DATA[party]["addresses"][0]
            
        party_details[party] = {
            "address": address,
            "contact": COMPANY_DATA[party]["contact"]
        }
    
    return party_details

def generate_contract_text(contract_type, parties, dates, version=None):
    """Generate contract text based on the parameters."""
    party_names = list(parties.keys())
    
    contract_text = f"""
CONTRACT: {contract_type.upper()}
{'VERSION ' + str(version) if version else ''}

EFFECTIVE DATE: {dates['creation'].strftime('%B %d, %Y')}
EXPIRATION DATE: {dates['expiration'].strftime('%B %d, %Y')}
RENEWAL DATE: {dates['renewal'].strftime('%B %d, %Y')}

BETWEEN:
{party_names[0]} (hereinafter referred to as "Party A")
Address: {parties[party_names[0]]['address']}
Contact: {parties[party_names[0]]['contact'][0]}, {parties[party_names[0]]['contact'][1]}

AND:
{party_names[1]} (hereinafter referred to as "Party B")
Address: {parties[party_names[1]]['address']}
Contact: {parties[party_names[1]]['contact'][0]}, {parties[party_names[1]]['contact'][1]}

TERMS AND CONDITIONS:

1. SCOPE OF WORK
   Party A agrees to provide {contract_type} services to Party B as described in Exhibit A.

2. TERM
   This agreement shall commence on {dates['creation'].strftime('%B %d, %Y')} and continue until {dates['expiration'].strftime('%B %d, %Y')}, 
   unless terminated earlier as provided herein.

3. RENEWAL
   This agreement may be renewed by written consent of both parties no later than {dates['renewal'].strftime('%B %d, %Y')}.

4. COMPENSATION
   Party B shall pay Party A according to the fee schedule outlined in Exhibit B.

5. TERMINATION
   Either party may terminate this agreement with 30 days written notice to the other party.

6. CONFIDENTIALITY
   Both parties agree to maintain the confidentiality of all information marked as confidential.

7. GOVERNING LAW
   This agreement shall be governed by the laws of the state of {random.choice(['New York', 'California', 'Illinois', 'Massachusetts'])}.

IN WITNESS WHEREOF, the parties have executed this agreement as of the Effective Date.

{party_names[0]}                                {party_names[1]}
____________________                       ____________________
Signature                                  Signature

Date: {dates['creation'].strftime('%B %d, %Y')}                    Date: {dates['creation'].strftime('%B %d, %Y')}
"""
    return contract_text

def create_txt_file(filename, content, metadata=None):
    """Create a text file with contract content."""
    try:
        with open(os.path.join(OUTPUT_DIR, filename), 'w') as f:
            if metadata:
                f.write("--- METADATA ---\n")
                for key, value in metadata.items():
                    f.write(f"{key}: {value}\n")
                f.write("---------------\n\n")
            
            f.write(content)
        logger.info(f"Created text file: {filename}")
    except Exception as e:
        logger.error(f"Failed to create text file {filename}: {str(e)}")

def create_docx_file(filename, content, metadata=None):
    """Create a Word document with contract content."""
    try:
        doc = Document()
        
        if metadata:
            doc.add_heading("METADATA", 1)
            for key, value in metadata.items():
                doc.add_paragraph(f"{key}: {value}")
            doc.add_paragraph("---------------")
            
        # Split content by lines and add each paragraph
        for para in content.split('\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        
        doc.save(os.path.join(OUTPUT_DIR, filename))
        logger.info(f"Created Word document: {filename}")
    except Exception as e:
        logger.error(f"Failed to create Word document {filename}: {str(e)}")

def create_pdf_file(filename, content, metadata=None, scanned=False):
    """Create a PDF with contract content."""
    try:
        import reportlab
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.units import inch
        
        packet = io.BytesIO()
        
        if scanned:
            # Create an image with text (simulating a scanned document)
            img = Image.new('RGB', (2100, 2970), color=(255, 255, 255))
            d = ImageDraw.Draw(img)
            
            # Try to load a font, or use default if not available
            try:
                font = ImageFont.truetype("Arial", 24)
            except:
                font = ImageFont.load_default()
                
            y_position = 100
            for line in content.split('\n'):
                if line.strip():
                    d.text((100, y_position), line.strip(), fill=(0, 0, 0), font=font)
                    y_position += 40
            
            # Save image as PDF
            img_path = os.path.join(OUTPUT_DIR, f"{filename.replace('.pdf', '')}_scan.png")
            img.save(img_path)
            
            # Convert image to PDF
            from reportlab.lib.utils import ImageReader
            c = canvas.Canvas(packet, pagesize=letter)
            c.drawImage(ImageReader(img_path), 0, 0, width=8.5*inch, height=11*inch)
            c.save()
            
            # Remove temporary image file
            os.remove(img_path)
            
        else:
            # Create normal PDF with text
            doc = SimpleDocTemplate(packet, pagesize=letter)
            styles = getSampleStyleSheet()
            flowables = []
            
            if metadata:
                flowables.append(Paragraph("<b>METADATA</b>", styles["Heading1"]))
                for key, value in metadata.items():
                    flowables.append(Paragraph(f"{key}: {value}", styles["Normal"]))
                flowables.append(Paragraph("---------------", styles["Normal"]))
            
            for para in content.split('\n'):
                if para.strip():
                    flowables.append(Paragraph(para.strip(), styles["Normal"]))
            
            doc.build(flowables)
        
        # Move to the beginning of the StringIO buffer
        packet.seek(0)
        
        # Create a new PDF with the content
        new_pdf = PdfReader(packet)
        output = PdfWriter()
        
        # Add page from the created PDF
        for page in new_pdf.pages:
            output.add_page(page)
        
        # Write the output PDF
        output_path = os.path.join(OUTPUT_DIR, filename)
        with open(output_path, 'wb') as f:
            output.write(f)
        
        logger.info(f"Created PDF document: {filename}" + (" (scanned)" if scanned else ""))
    except Exception as e:
        logger.error(f"Failed to create PDF {filename}: {str(e)}")

def create_unstructured_text(filename, related_contract, meeting_notes=True):
    """Create unstructured text like meeting notes or email correspondence."""
    try:
        contract_type = related_contract.split('_')[0]
        
        if meeting_notes:
            content = f"""
Meeting Notes: Discussion about {contract_type} Contract
Date: {datetime.now().strftime('%B %d, %Y')}

Attendees:
- Sarah Johnson (Legal)
- Michael Chen (Procurement)
- David Rodriguez ({random.choice(DEPARTMENTS)})
- Jennifer Smith (Stakeholder)

Notes:
1. Reviewed the current {contract_type} contract with {random.choice(list(COMPANY_DATA.keys()))}
2. Concerns about section 3.4 regarding liability coverage - needs revision
3. The renewal date seems to be {"approaching" if random.random() < 0.5 else "far off"}, but we should prepare documentation
4. {random.choice(list(COMPANY_DATA.keys()))} has {random.choice(["changed their address recently", "a new point of contact", "requested amendments"])}
5. Action items:
   a. David to follow up on pricing discrepancies
   b. Sarah to review legal implications of clause 7.2
   c. Jennifer to confirm {random.choice(["delivery schedule", "payment terms", "scope of work"])}

Next steps: Schedule follow-up meeting in two weeks to finalize the {'contract amendments' if random.random() < 0.5 else 'renewal terms'}
"""
        else:
            # Email correspondence
            content = f"""
From: {random.choice(['sarah.johnson@company.com', 'michael.chen@company.com', 'david.rodriguez@company.com'])}
To: {random.choice(['legal@company.com', 'procurement@company.com', 'contracts@company.com'])}
Subject: {contract_type} Contract with {random.choice(list(COMPANY_DATA.keys()))}
Date: {datetime.now().strftime('%B %d, %Y %H:%M:%S')}

Hi team,

I wanted to follow up regarding the {contract_type} contract that we discussed last week. 
There are a few points that need attention:

1. The contract {random.choice(['expires on', 'needs renewal by', 'has issues with'])} the date mentioned in section 2.1.
2. We need to clarify the {random.choice(['payment terms', 'deliverables', 'termination clause'])}.
3. {random.choice(list(COMPANY_DATA.keys()))} has requested {random.choice(['a change in terms', 'an extension', 'a discount'])}.

Can we schedule a call to discuss these matters? I'm available {random.choice(['tomorrow afternoon', 'next Monday', 'later this week'])}.

Best regards,
{random.choice(['Sarah', 'Michael', 'David'])}
"""
        
        with open(os.path.join(OUTPUT_DIR, filename), 'w') as f:
            f.write(content)
        logger.info(f"Created unstructured text: {filename}")
    except Exception as e:
        logger.error(f"Failed to create unstructured text {filename}: {str(e)}")

def generate_dataset():
    """Generate the complete synthetic dataset."""
    logger.info("Starting synthetic data generation")
    
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    # Generate contract versions with conflicts
    contract_versions = []
    
    # Generate base contracts (will create multiple versions of some)
    base_contracts = []
    for i in range(6):
        contract_type = random.choice(CONTRACT_TYPES)
        department = random.choice(DEPARTMENTS)
        
        # Decide if this contract will have multiple versions
        has_multiple_versions = random.random() < 0.3
        
        # Base contract info
        contract_info = {
            "type": contract_type,
            "department": department,
            "has_versions": has_multiple_versions,
            "has_conflict": random.random() < 0.4,  # 40% chance of conflicts
            "has_metadata": random.random() < 0.7,  # 70% chance of having metadata
        }
        base_contracts.append(contract_info)
    
    logger.info(f"Generated {len(base_contracts)} base contract configurations")
    
    # Create all document formats based on the base contracts
    documents_created = 0
    
    # Track expiration dates for testing
    upcoming_expirations = []
    
    for contract in base_contracts:
        contract_type = contract["type"]
        department = contract["department"]
        base_filename = f"{contract_type.replace(' ', '')}_Contract"
        
        # Generate dates and parties
        dates = generate_contract_dates(with_conflicts=contract["has_conflict"])
        parties = select_parties(conflict=contract["has_conflict"])
        
        # Check if this contract will expire soon (for testing)
        days_to_expiration = (dates['expiration'] - CURRENT_DATE).days
        if days_to_expiration < 30 and days_to_expiration > 0:
            upcoming_expirations.append({
                "name": base_filename,
                "expiration": dates['expiration'],
                "days_remaining": days_to_expiration
            })
        
        # Create metadata if needed
        metadata = None
        if contract["has_metadata"]:
            metadata = {
                "Contract Name": f"{contract_type} - {list(parties.keys())[0]} and {list(parties.keys())[1]}",
                "Department": department,
                "Contract Owner": f"{random.choice(['John Smith', 'Jane Doe', 'Robert Johnson', 'Emily Williams'])}",
                "Contract ID": f"CNT-{random.randint(10000, 99999)}"
            }
        
        # Generate base contract text
        contract_text = generate_contract_text(contract_type, parties, dates)
        
        # Determine what formats to create for this contract
        formats = random.sample([
            ("txt", create_txt_file), 
            ("docx", create_docx_file), 
            ("pdf", lambda f, c, m: create_pdf_file(f, c, m, False)),
            ("pdf_scanned", lambda f, c, m: create_pdf_file(f, c, m, True))
        ], random.randint(1, 3))
        
        # Create the primary contract document
        for format_info in formats:
            fmt, create_func = format_info
            if fmt == "pdf_scanned":
                filename = f"{base_filename}_scanned.pdf"
            else:
                filename = f"{base_filename}.{fmt}"
            
            create_func(filename, contract_text, metadata)
            documents_created += 1
        
        # Create contract versions if needed
        if contract["has_versions"]:
            # Generate 1-2 additional versions
            for version in range(2, 2 + random.randint(1, 2)):
                # Generate slightly modified contract with potential conflicts
                modified_dates = generate_contract_dates(with_conflicts=True)
                modified_parties = select_parties(conflict=True)
                
                # Modify metadata for the new version
                if metadata:
                    modified_metadata = metadata.copy()
                    modified_metadata["Version"] = str(version)
                else:
                    modified_metadata = {"Version": str(version)}
                
                # Generate modified contract text
                modified_text = generate_contract_text(
                    contract_type, 
                    modified_parties, 
                    modified_dates, 
                    version=version
                )
                
                # Create the new version in a random format
                format_info = random.choice(formats)
                fmt, create_func = format_info
                
                if fmt == "pdf_scanned":
                    version_filename = f"{base_filename}_v{version}_scanned.pdf"
                else:
                    version_filename = f"{base_filename}_v{version}.{fmt}"
                
                create_func(version_filename, modified_text, modified_metadata)
                documents_created += 1
        
        # Create related unstructured documents
        if random.random() < 0.4:  # 40% chance of having meeting notes
            create_unstructured_text(
                f"{base_filename}_meeting_notes.txt", 
                base_filename, 
                meeting_notes=True
            )
            documents_created += 1
        
        if random.random() < 0.4:  # 40% chance of having email correspondence
            create_unstructured_text(
                f"{base_filename}_correspondence.txt", 
                base_filename, 
                meeting_notes=False
            )
            documents_created += 1
    
    logger.info(f"Created {documents_created} documents in total")
    logger.info(f"Contracts with upcoming expiration: {len(upcoming_expirations)}")
    
    # Add at least one contract with upcoming expiration if none were created randomly
    if not upcoming_expirations:
        # Create a contract that expires soon
        contract_type = random.choice(CONTRACT_TYPES)
        department = random.choice(DEPARTMENTS)
        base_filename = f"{contract_type.replace(' ', '')}_UrgentExpiration"
        
        # Generate dates with upcoming expiration
        creation_date = CURRENT_DATE - timedelta(days=360)
        expiration_date = CURRENT_DATE + timedelta(days=random.randint(5, 25))
        renewal_date = expiration_date - timedelta(days=30)
        
        dates = {
            "creation": creation_date,
            "expiration": expiration_date,
            "renewal": renewal_date
        }
        
        parties = select_parties(conflict=False)
        contract_text = generate_contract_text(contract_type, parties, dates)
        
        # Add metadata
        metadata = {
            "Contract Name": f"{contract_type} - {list(parties.keys())[0]} and {list(parties.keys())[1]}",
            "Department": department,
            "Contract Owner": f"{random.choice(['John Smith', 'Jane Doe', 'Robert Johnson', 'Emily Williams'])}",
            "Contract ID": f"CNT-{random.randint(10000, 99999)}",
            "Urgent": "Yes - Expiring Soon"
        }
        
        # Create as PDF
        create_pdf_file(f"{base_filename}.pdf", contract_text, metadata)
        documents_created += 1
        
        upcoming_expirations.append({
            "name": base_filename,
            "expiration": expiration_date,
            "days_remaining": (expiration_date - CURRENT_DATE).days
        })
        
        logger.info(f"Added urgent expiring contract: {base_filename}.pdf")
    
    logger.info(f"Dataset generation complete. {documents_created} documents created.")
    return documents_created

if __name__ == "__main__":
    num_documents = generate_dataset()
    print(f"Generated {num_documents} synthetic contract documents in {OUTPUT_DIR}")